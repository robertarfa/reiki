import pandas as pd
from fpdf import FPDF
from docx import Document


df = pd.read_excel('../data_raw/AvaChakras.xlsx')

df.head()
df.columns
varId = "Id"

def validate_id():
    if(df.columns[0] == 'ID'):
        varId = "ID"
    return varId
        
def filter_by_id(new_id):    
    # Find columns that start with "Pontos – ["
    relevant_columns = [col for col in df.columns if (col.startswith("Pontos – [") | col.startswith("Points - ["))]

    # Group by 'E-mail:' and sum the relevant columns
    grouped_df = df.groupby(varId)[relevant_columns].sum()
    # Reset index to make 'E-mail:' a regular column
    grouped_df = grouped_df.reset_index()

    # Melt the DataFrame to long format
    melted_df = grouped_df.melt(id_vars=varId, var_name='Chakra', value_name='Soma')

    # Merge with the data for Id == 1
    return melted_df[melted_df[varId] == new_id]

def dados_resumidos(id):
    """Retorna um dicionário com dados resumidos para um ID específico, ou None se o ID não existir."""
    linha = df[df[varId] == id]  # Filtra o DataFrame
    if linha.empty:
        return None  # Retorna None se nenhuma linha corresponder ao ID
    else:
        linha = linha.iloc[0]  # Acessa a primeira (e única) linha
        return {
            "Id": linha[varId],
            "Nome": linha["Nome Completo:"],
            "E-mail": linha["E-mail:"],
            "Whatsapp": linha["""Whatsapp:
(Exemplo: 5511988881111)"""]
        }

def pontuacao_chakras(id_filtered):
    resultados = {}
    for chakra in chakras:
        soma = id_filtered[id_filtered['Chakra'].str.contains(chakra)]['Soma'].sum()
        resultados[chakra] = soma

    dados_corrigidos = dict(map(lambda item: (item[0].replace('\\', '').replace('[', '').replace(']', ''), item[1]), resultados.items()))
    tabela = pd.DataFrame(list(dados_corrigidos.items()), columns=['Chakra', 'Pontos'])

    return tabela.sort_values(by='Pontos', ascending=True)

def chakra_menor_pontuacao(tabela_ordenada):
    min_nota = tabela_ordenada['Pontos'].min()
    # Seleciona as linhas com o valor mínimo
    min_rows = tabela_ordenada[tabela_ordenada['Pontos'] == min_nota]
    # Imprime as linhas com o valor mínimo

    return min_rows    

chakras = [r'\[SEXUAL]', r'\[CARDÍACO]', r'\[SOLAR]',
           r'\[FRONTAL]', r'\[CORONÁRIO]', r'\[BASE]',
           r'\[LARÍNGEO]']

id = 3
varId = validate_id()
dados_filtrados = filter_by_id(3)
dados = dados_resumidos(3)

print("Nome: ", dados["Nome"])
print("E-mail: ", dados["E-mail"])
print("Whatsapp: ", dados["Whatsapp"])
print("\n")

pontuacao_chakras = pontuacao_chakras(dados_filtrados)
chakra_menor_pontuacao = chakra_menor_pontuacao(pontuacao_chakras)

print("Pontuação dos Chakras: ")
print(pontuacao_chakras)
print("\n")

print("Chakras Com menos pontos: ")
print(chakra_menor_pontuacao)
print("\n")


Frases = []
for chakra in chakra_menor_pontuacao["Chakra"]:
    soma = dados_filtrados[(dados_filtrados['Chakra'].str.contains(chakra)) & (dados_filtrados['Soma'] == 0)]["Chakra"]
    Frases.append(soma)

# tabela_final = pd.concat(Frases).reset_index()

# tabela_final['Chakra'] = tabela_final['Chakra'].str.replace('Pontos – ', '', regex=False)

print("Respostas erradas: ")
print(Frases)

dados_pdf = {
    "Informações Pessoais": {
        "Nome": dados["Nome"],
        "E-mail": dados["E-mail"],
        "Whatsapp": dados["Whatsapp"]
    },
    "Pontuação dos Chakras": pontuacao_chakras ,
    "Chakras em Desequilíbrio": chakra_menor_pontuacao,
    "Respostas Erradas": None # Initialize to None
}

#Improved handling of Respostas Erradas
if not chakra_menor_pontuacao.empty: #Check if chakra_menor_pontuacao is not empty before proceeding.
    respostas_erradas = pd.concat([
        dados_filtrados[(dados_filtrados['Chakra'].str.contains(chakra)) & (dados_filtrados['Soma'] == 0)][['Chakra']] #Select only the 'Chakra' column
        for chakra in chakra_menor_pontuacao["Chakra"]
    ], ignore_index=True) #add ignore_index=True to avoid issues with duplicate indices

    respostas_erradas["Chakra"] = respostas_erradas["Chakra"].str.replace('Pontos – ', ' - ', regex=False).str.replace('Points - ', ' - ', regex=False).replace('"', '', regex=False)
    if not respostas_erradas.empty: #Check if any wrong answers exist after concatenation.
        dados_pdf["Respostas Erradas"] = respostas_erradas

# #Criando o PDF
# pdf = FPDF()
# pdf.add_page()
# pdf.set_font("Arial", size=12)

# for secao, conteudo in dados_pdf.items():
#     pdf.set_font("Arial", "B", 12) # Define fonte em negrito (Arial, tamanho 12)
#     pdf.cell(200, 10, txt=f"{secao}", ln=1, align="C", border=1)
#     pdf.set_font("Arial", "", 12) # Volta para fonte normal (opcional, dependendo do resto do seu documento)

#     if isinstance(conteudo, dict):
#         for chave, valor in conteudo.items():
#             pdf.cell(200, 10, txt=f"{chave}: {valor}", ln=1, align="L")
#     elif isinstance(conteudo, pd.DataFrame):
#         buffer = io.StringIO()
#         conteudo.to_csv(buffer, index=False, header=False, encoding='utf-8')
#         buffer.seek(0)
#         for line in buffer.readlines():
#             cleaned_line = line.strip().replace('\u2013', '-')
#             pdf.multi_cell(0, 10, txt=cleaned_line, align="L")
#         buffer.close()
#     elif conteudo is None:
#         pdf.multi_cell(0, 10, txt="Nenhuma resposta errada encontrada.", align="L")


# pdf.output("relatorio.pdf")


# Criando o documento do Word
document = Document()

def add_heading(document, text, level=1):
    document.add_heading(text, level)

def add_paragraph(document, text):
    document.add_paragraph(text)


def add_table(document, dataframe):
    table = document.add_table(rows=1, cols=len(dataframe.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(dataframe.columns):
        hdr_cells[i].text = col

    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)


for secao, conteudo in dados_pdf.items():
    add_heading(document, f"{secao}:", level=1)
    if isinstance(conteudo, dict):
        for chave, valor in conteudo.items():
            add_paragraph(document, f"{chave}: {valor}")
    elif isinstance(conteudo, pd.DataFrame):
        if not conteudo.empty:
            add_table(document, conteudo)
        else:
            add_paragraph(document, "Nenhuma resposta errada encontrada.")
    elif conteudo is None:
        add_paragraph(document, "Nenhuma resposta errada encontrada.")

import datetime

hoje = datetime.date.today()
print(f"hoje: {hoje}")

nome = f"{dados['Nome'].strip().replace('/', '_').replace('\\', '_').replace("  ", " ")}"
email = f"{dados['E-mail'].strip().replace('/', '_').replace('\\', '_').replace("  ", " ")}"
full = f"{nome}-{email}"
filename = f"{dados["Id"]}-{full}-{hoje.strftime('%d%m%y')}.docx"
path = f"../documentos/{filename}"

document.save(path)